from docx import Document
from datetime import datetime, timedelta
import imaplib
import email
from email.header import decode_header
import requests
import os
import re
import time
from flask import Flask, render_template, jsonify, request, redirect, session, flash, send_file
import logging
import sqlite3
import json
import traceback
import threading
import bcrypt
from functools import wraps
import io
import csv
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from fpdf import FPDF

# Initialize Flask app
app = Flask(__name__)

# Generate a secure secret key (set this in Railway environment variables)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'dev-secret-key-change-in-production')

# Default admin credentials (set these in Railway environment variables)
DEFAULT_USERNAME = os.getenv('DASHBOARD_USERNAME', 'admin')
DEFAULT_PASSWORD = os.getenv('DASHBOARD_PASSWORD', 'admin123')

# Hash the password
HASHED_PASSWORD = bcrypt.hashpw(DEFAULT_PASSWORD.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
print(f"🔐 Password hash initialized for user: {DEFAULT_USERNAME}")

def get_db_path():
    """Get database path that works with Railway Volume"""
    # Use Railway volume path if available, otherwise use local path
    db_path = os.getenv('DATABASE_PATH', '/data/email_summaries.db')
    
    # Ensure the directory exists
    db_dir = os.path.dirname(db_path)
    if db_dir and not os.path.exists(db_dir):
        os.makedirs(db_dir, exist_ok=True)
        print(f"📁 Created database directory: {db_dir}")
    
    print(f"📁 Using database at: {db_path}")
    return db_path

# ==================== AUTHENTICATION DECORATORS ====================

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'logged_in' not in session:
            return redirect('/login?next=' + request.path)
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'logged_in' not in session or session.get('username') != DEFAULT_USERNAME:
            return redirect('/login?next=' + request.path)
        return f(*args, **kwargs)
    return decorated_function

# ==================== AUTHENTICATION ROUTES ====================

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login page"""
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()
        next_page = request.args.get('next', '/dashboard')
        
        # Verify credentials
        if username == DEFAULT_USERNAME:
            try:
                if bcrypt.checkpw(password.encode('utf-8'), HASHED_PASSWORD.encode('utf-8')):
                    session['logged_in'] = True
                    session['username'] = username
                    session['login_time'] = datetime.now().isoformat()
                    session.permanent = True
                    app.permanent_session_lifetime = timedelta(hours=24)
                    print(f"🔐 User '{username}' logged in successfully")
                    return redirect(next_page)
            except Exception as e:
                print(f"🔐 Login error: {e}")
        
        flash('Invalid username or password', 'error')
        return render_template('login.html')
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    """Logout user"""
    username = session.get('username', 'Unknown')
    session.clear()
    print(f"🔐 User '{username}' logged out")
    flash('You have been logged out successfully', 'success')
    return redirect('/login')

@app.route('/change-password', methods=['GET', 'POST'])
@admin_required
def change_password():
    """Change password page (admin only)"""
    global HASHED_PASSWORD
    
    if request.method == 'POST':
        current_password = request.form.get('current_password', '').strip()
        new_password = request.form.get('new_password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()
        
        # Validate current password
        try:
            if not bcrypt.checkpw(current_password.encode('utf-8'), HASHED_PASSWORD.encode('utf-8')):
                flash('Current password is incorrect', 'error')
                return render_template('change_password.html')
        except Exception as e:
            flash('Error validating current password', 'error')
            return render_template('change_password.html')
        
        # Validate new password
        if new_password != confirm_password:
            flash('New passwords do not match', 'error')
            return render_template('change_password.html')
        
        if len(new_password) < 8:
            flash('New password must be at least 8 characters', 'error')
            return render_template('change_password.html')
        
        # Update password
        try:
            HASHED_PASSWORD = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        except Exception as e:
            flash('Error hashing new password', 'error')
            return render_template('change_password.html')
        
        session.clear()
        print("🔐 Password changed successfully")
        flash('Password changed successfully. Please log in again.', 'success')
        return redirect('/login')
    
    return render_template('change_password.html')

# ==================== DATABASE INITIALIZATION ====================

def init_db():
    """Initialize SQLite database"""
    db_path = get_db_path()
    print(f"📁 Initializing database at: {db_path}")
    
    try:
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        # Create tables
        c.execute('''
            CREATE TABLE IF NOT EXISTS summary_runs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                run_date TEXT NOT NULL,
                total_emails INTEGER,
                processed_emails INTEGER,
                success_rate REAL,
                deepseek_tokens INTEGER,
                status TEXT,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        c.execute('''
            CREATE TABLE IF NOT EXISTS email_data (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                run_id INTEGER,
                email_number INTEGER,
                sender TEXT,
                receiver TEXT,
                subject TEXT,
                summary TEXT,
                email_date TEXT,
                processed_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (run_id) REFERENCES summary_runs (id)
            )
        ''')
        
        # Create indexes
        c.execute('CREATE INDEX IF NOT EXISTS idx_run_id ON email_data (run_id)')
        c.execute('CREATE INDEX IF NOT EXISTS idx_email_number ON email_data (email_number)')
        c.execute('CREATE INDEX IF NOT EXISTS idx_email_date ON email_data (email_date)')
        
        conn.commit()
        conn.close()
        print(f"✅ Database initialized successfully at: {db_path}")
        return True
    except Exception as e:
        print(f"❌ Database initialization error: {e}")
        return False

# Initialize database
init_db()

# ==================== EMAIL PROCESSING FUNCTIONS ====================

def fetch_and_summarize_emails():
    """Main function to fetch emails and generate summaries"""
    print("🚀 Starting email fetch and summarization process...")
    
    try:
        # Get email credentials from environment variables
        email_account = os.getenv('EMAIL_ACCOUNT')
        email_password = os.getenv('EMAIL_PASSWORD')
        imap_server = os.getenv('IMAP_SERVER', 'imap.gmail.com')
        deepseek_api_key = os.getenv('DEEPSEEK_API_KEY')
        
        if not all([email_account, email_password, deepseek_api_key]):
            print("❌ Missing required environment variables")
            return False
        
        # Connect to email server
        print(f"📧 Connecting to {imap_server}...")
        mail = imaplib.IMAP4_SSL(imap_server)
        mail.login(email_account, email_password)
        mail.select('inbox')
        
        # Calculate date for last 24 hours
        date_since = (datetime.now() - timedelta(days=1)).strftime('%d-%b-%Y')
        
        # Search for emails from last 24 hours
        result, data = mail.search(None, f'(SINCE "{date_since}")')
        email_ids = data[0].split()
        
        print(f"📧 Found {len(email_ids)} emails in the last 24 hours")
        
        if not email_ids:
            print("📧 No new emails found")
            mail.close()
            mail.logout()
            return True
        
        # Process emails
        processed_emails = []
        total_tokens = 0
        
        for i, email_id in enumerate(email_ids[:20]):  # Limit to 20 emails per run
            try:
                result, msg_data = mail.fetch(email_id, '(RFC822)')
                msg = email.message_from_bytes(msg_data[0][1])
                
                # Extract email details
                subject, encoding = decode_header(msg['Subject'])[0]
                if isinstance(subject, bytes):
                    subject = subject.decode(encoding if encoding else 'utf-8')
                
                from_addr = msg.get('From')
                to_addr = msg.get('To')
                date = msg.get('Date')
                
                # Get email body
                body = ""
                if msg.is_multipart():
                    for part in msg.walk():
                        if part.get_content_type() == 'text/plain':
                            body = part.get_payload(decode=True).decode()
                            break
                else:
                    body = msg.get_payload(decode=True).decode()
                
                # Clean body (remove extra whitespace, etc.)
                body = re.sub(r'\s+', ' ', body).strip()[:2000]  # Limit to 2000 chars
                
                # Generate summary using DeepSeek API
                summary = generate_summary(body, subject, deepseek_api_key)
                total_tokens += len(summary.split())  # Approximate token count
                
                processed_emails.append({
                    'number': i + 1,
                    'sender': from_addr,
                    'receiver': to_addr,
                    'subject': subject,
                    'summary': summary,
                    'date': date
                })
                
                print(f"✅ Processed email {i+1}: {subject[:50]}...")
                
            except Exception as e:
                print(f"❌ Error processing email {i+1}: {e}")
                continue
        
        # Save to database
        save_to_database(processed_emails, total_tokens)
        
        mail.close()
        mail.logout()
        
        print(f"✅ Email processing completed. Processed {len(processed_emails)} emails")
        return True
        
    except Exception as e:
        print(f"❌ Error in fetch_and_summarize_emails: {e}")
        return False

def generate_summary(body, subject, api_key):
    """Generate summary using DeepSeek API"""
    try:
        headers = {
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json'
        }
        
        prompt = f"Summarize this email in 2-3 sentences:\n\nSubject: {subject}\n\nBody: {body}"
        
        data = {
            'model': 'deepseek-chat',
            'messages': [{'role': 'user', 'content': prompt}],
            'max_tokens': 150
        }
        
        response = requests.post(
            'https://api.deepseek.com/v1/chat/completions',
            headers=headers,
            json=data,
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content'].strip()
        else:
            print(f"❌ DeepSeek API error: {response.status_code}")
            return "Summary generation failed"
            
    except Exception as e:
        print(f"❌ Error generating summary: {e}")
        return "Summary generation failed"

def save_to_database(emails, tokens_used):
    """Save processed emails to database"""
    try:
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        # Create summary run record
        run_date = datetime.now().isoformat()
        c.execute('''
            INSERT INTO summary_runs (run_date, total_emails, processed_emails, deepseek_tokens, status)
            VALUES (?, ?, ?, ?, ?)
        ''', (run_date, len(emails), len(emails), tokens_used, 'completed'))
        
        run_id = c.lastrowid
        
        # Save individual emails
        for email in emails:
            c.execute('''
                INSERT INTO email_data (run_id, email_number, sender, receiver, subject, summary, email_date)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (run_id, email['number'], email['sender'], email['receiver'], 
                  email['subject'], email['summary'], email['date']))
        
        conn.commit()
        conn.close()
        
        print(f"✅ Saved {len(emails)} emails to database (Run ID: {run_id})")
        
    except Exception as e:
        print(f"❌ Error saving to database: {e}")

# ==================== API ENDPOINTS ====================

@app.route('/api/trigger-summary', methods=['POST'])
@admin_required
def trigger_summary():
    """Manually trigger email summarization"""
    try:
        # Run in background thread to not block response
        thread = threading.Thread(target=fetch_and_summarize_emails)
        thread.start()
        
        return jsonify({
            'success': True,
            'message': 'Email summarization started in background'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/latest-summary', methods=['GET'])
@login_required
def get_latest_summary():
    """Get the latest email summary data"""
    try:
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        # Get latest run
        c.execute('SELECT * FROM summary_runs ORDER BY id DESC LIMIT 1')
        latest_run = c.fetchone()
        
        if not latest_run:
            return jsonify({
                'success': True,
                'data': {
                    'run_info': None,
                    'emails': []
                }
            })
        
        # Get emails from latest run
        c.execute('''
            SELECT email_number, sender, receiver, subject, summary, email_date 
            FROM email_data 
            WHERE run_id = ? 
            ORDER BY email_number
        ''', (latest_run[0],))
        
        emails = []
        for row in c.fetchall():
            emails.append({
                'number': row[0],
                'from': row[1],
                'to': row[2],
                'subject': row[3],
                'summary': row[4],
                'date': row[5]
            })
        
        conn.close()
        
        run_info = {
            'id': latest_run[0],
            'date': latest_run[1],
            'total': latest_run[2],
            'processed': latest_run[3],
            'tokens': latest_run[5],
            'status': latest_run[6]
        }
        
        return jsonify({
            'success': True,
            'data': {
                'run_info': run_info,
                'emails': emails
            }
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/stats', methods=['GET'])
@login_required
def get_stats():
    """Get dashboard statistics"""
    try:
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        # Total emails processed
        c.execute('SELECT COUNT(*) FROM email_data')
        total_emails = c.fetchone()[0]
        
        # Total runs
        c.execute('SELECT COUNT(*) FROM summary_runs')
        total_runs = c.fetchone()[0]
        
        # Last run date
        c.execute('SELECT run_date FROM summary_runs ORDER BY id DESC LIMIT 1')
        last_run = c.fetchone()
        last_run_date = last_run[0] if last_run else None
        
        # Average tokens per run
        c.execute('SELECT AVG(deepseek_tokens) FROM summary_runs')
        avg_tokens = c.fetchone()[0] or 0
        
        conn.close()
        
        return jsonify({
            'success': True,
            'data': {
                'total_emails': total_emails,
                'total_runs': total_runs,
                'last_run': last_run_date,
                'avg_tokens': int(avg_tokens)
            }
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ==================== EXPORT ROUTES ====================

@app.route('/api/export-data', methods=['POST'])
@login_required
def export_data():
    """Export email data in various formats"""
    try:
        data = request.json
        format_type = data.get('format', 'csv')
        filename = data.get('filename', 'email_summaries')
        
        # Get filters from request
        filters = data.get('filters', {})
        
        # Get email data with filters
        email_data = get_filtered_email_data(filters)
        
        if not email_data:
            return jsonify({"error": "No data to export"}), 400
        
        if format_type == 'csv':
            return export_csv(email_data, filename)
        elif format_type == 'json':
            return export_json(email_data, filename)
        elif format_type == 'word':
            return export_word(email_data, filename)
        elif format_type == 'pdf':
            return export_pdf(email_data, filename)
        else:
            return jsonify({"error": "Unsupported format"}), 400
            
    except Exception as e:
        print(f"❌ Export error: {e}")
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

def get_filtered_email_data(filters):
    """Get email data with optional filters"""
    try:
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        # Build query based on filters
        query = '''
            SELECT email_number, sender, receiver, subject, summary, email_date 
            FROM email_data 
            WHERE 1=1
        '''
        params = []
        
        # Apply date filter
        date_range = filters.get('dateRange')
        if date_range == 'today':
            today = datetime.now().strftime('%Y-%m-%d')
            query += " AND DATE(email_date) = ?"
            params.append(today)
        elif date_range == 'week':
            week_ago = (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%d')
            query += " AND DATE(email_date) >= ?"
            params.append(week_ago)
        elif date_range == 'month':
            month_ago = (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d')
            query += " AND DATE(email_date) >= ?"
            params.append(month_ago)
        elif date_range == 'custom' and filters.get('startDate') and filters.get('endDate'):
            query += " AND DATE(email_date) BETWEEN ? AND ?"
            params.extend([filters['startDate'], filters['endDate']])
        
        # Apply sender/receiver filters
        sender = filters.get('sender')
        if sender:
            query += " AND sender LIKE ?"
            params.append(f'%{sender}%')
        
        receiver = filters.get('receiver')
        if receiver:
            query += " AND receiver LIKE ?"
            params.append(f'%{receiver}%')
        
        # Apply search query
        search = filters.get('search')
        if search:
            query += " AND (sender LIKE ? OR receiver LIKE ? OR subject LIKE ? OR summary LIKE ?)"
            params.extend([f'%{search}%', f'%{search}%', f'%{search}%', f'%{search}%'])
        
        # Get latest run if no specific date range
        if not date_range and not filters.get('allRuns', False):
            c.execute('SELECT id FROM summary_runs ORDER BY id DESC LIMIT 1')
            latest_run = c.fetchone()
            if latest_run:
                query += " AND run_id = ?"
                params.append(latest_run[0])
        
        query += " ORDER BY email_number"
        
        c.execute(query, params)
        rows = c.fetchall()
        
        email_data = []
        for row in rows:
            email_data.append({
                "number": row[0],
                "from": row[1],
                "to": row[2],
                "subject": row[3],
                "summary": row[4],
                "date": row[5]
            })
        
        conn.close()
        return email_data
        
    except Exception as e:
        print(f"❌ Error getting filtered email data: {e}")
        return []

def export_csv(email_data, filename):
    """Export email data as CSV"""
    try:
        output = io.StringIO()
        writer = csv.writer(output)
        
        writer.writerow(['#', 'Sender', 'Receiver', 'Date', 'Subject', 'Summary'])
        
        for email in email_data:
            writer.writerow([
                email['number'],
                email['from'],
                email['to'],
                email.get('date', ''),
                email['subject'],
                email['summary']
            ])
        
        output.seek(0)
        mem_file = io.BytesIO()
        mem_file.write(output.getvalue().encode('utf-8'))
        mem_file.seek(0)
        
        return send_file(
            mem_file,
            mimetype='text/csv',
            as_attachment=True,
            download_name=f'{filename}.csv'
        )
    except Exception as e:
        print(f"❌ CSV export error: {e}")
        raise

def export_json(email_data, filename):
    """Export email data as JSON"""
    try:
        mem_file = io.BytesIO()
        mem_file.write(json.dumps(email_data, indent=2).encode('utf-8'))
        mem_file.seek(0)
        
        return send_file(
            mem_file,
            mimetype='application/json',
            as_attachment=True,
            download_name=f'{filename}.json'
        )
    except Exception as e:
        print(f"❌ JSON export error: {e}")
        raise

def export_word(email_data, filename):
    """Export email data as Word document"""
    try:
        doc = Document()
        
        title = doc.add_heading('Email Summary Report', 0)
        title.alignment = 1
        
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Emails: {len(email_data)}")
        doc.add_paragraph()
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ['#', 'Sender', 'Receiver', 'Subject', 'Summary']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for email in email_data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(email['number'])
            row_cells[1].text = str(email['from'])[:40]
            row_cells[2].text = str(email['to'])[:40]
            row_cells[3].text = str(email['subject'])[:80]
            row_cells[4].text = str(email['summary'])[:400]
        
        mem_file = io.BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)
        
        return send_file(
            mem_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{filename}.docx'
        )
    except Exception as e:
        print(f"❌ Word export error: {e}")
        raise

def export_pdf(email_data, filename):
    """Export email data as PDF"""
    try:
        mem_file = io.BytesIO()
        doc = SimpleDocTemplate(mem_file, pagesize=letter)
        elements = []
        
        styles = getSampleStyleSheet()
        title = Paragraph("Email Summary Report", styles['Title'])
        elements.append(title)
        
        elements.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        elements.append(Paragraph(f"Total Emails: {len(email_data)}", styles['Normal']))
        elements.append(Spacer(1, 12))
        
        table_data = [['#', 'Sender', 'Receiver', 'Subject', 'Summary']]
        
        for email in email_data:
            table_data.append([
                str(email['number']),
                str(email['from'])[:30],
                str(email['to'])[:30],
                str(email['subject'])[:40],
                str(email['summary'])[:100]
            ])
        
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(table)
        doc.build(elements)
        
        mem_file.seek(0)
        
        return send_file(
            mem_file,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{filename}.pdf'
        )
    except Exception as e:
        print(f"❌ PDF export error: {e}")
        raise

# ==================== WEB ROUTES ====================

@app.route('/')
def index():
    """Redirect to dashboard if logged in, otherwise to login"""
    if 'logged_in' in session:
        return redirect('/dashboard')
    return redirect('/login')

@app.route('/dashboard')
@login_required
def dashboard():
    """Main dashboard page"""
    return render_template('dashboard.html', username=session.get('username'))

@app.route('/health')
def health():
    """Health check endpoint for Railway"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'database': get_db_path()
    })

# ==================== MAIN ENTRY POINT ====================

if __name__ == '__main__':
    port = int(os.getenv('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
